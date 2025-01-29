from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QTableWidget, QTableWidgetItem, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QFileDialog, QCheckBox, QMessageBox, QHeaderView, QProgressBar, QDialog, QTextBrowser
from PyQt5.QtGui import QTextCursor, QFont, QTextCharFormat, QColor, QBrush, QTextDocument
from PyQt5.QtCore import Qt, QRect, QThread, pyqtSignal, QRegularExpression, pyqtSlot, QTimer, QMetaObject, Q_ARG
from docx import Document
from docx.text.paragraph import Paragraph
import os
import win32com.client as win32
import re
import concurrent.futures
from whoosh.index import create_in, open_dir
from whoosh.fields import Schema, TEXT, ID, NUMERIC
from whoosh.qparser import QueryParser, RegexPlugin

class FindAllDocXFiles(QThread):
    progress_update = pyqtSignal(int, int)
    finished = pyqtSignal(list)

    def __init__(self, parent, dir_path):
        super().__init__()
        self.parent = parent
        self.dir_path = dir_path
        self.word_docs = []
        self.cancelled = False
    
    def run(self):
        # Search for Word documents in directory
        num_files = 0
        try:
            for root, dirs, files in os.walk(self.dir_path):
                for file in files:
                    if self.cancelled:
                        return
                    if file.endswith('.docx'):
                        self.word_docs.append(os.path.join(root, file))
                        num_files += 1
                        self.progress_update.emit(0, num_files)
        except Exception as e:
            message_box = QMessageBox(self.parent)
            message_box.setWindowTitle("Error while finding docx files.")
            message_box.setText(e)
            message_box.setStandardButtons(QMessageBox.Ok)
            message_box.setDefaultButton(QMessageBox.Ok)
            message_box.exec_()
            self.finished.emit()
            return
        
        self.finished.emit(self.word_docs)
    
    def cancel(self):
        self.cancelled = True

class SearchWordsThread(QThread):
    progress_update = pyqtSignal(int, int)
    finished = pyqtSignal()

    def __init__(self, parent, search_term, word_search_option):
        super().__init__()
        self.parent = parent
        self.search_term = search_term
        self.word_search_option = word_search_option
        self.cancelled = False
        self.executor = concurrent.futures.ThreadPoolExecutor()
        self.word_docs = []
        self.results = []
    
    def add_file_list(self, file_list):
        self.word_docs = file_list
    
    def add_file_list_and_run(self, file_list):
        self.add_file_list(file_list)
        self.run()
    
    def run(self):
        with self.parent.ix.searcher() as searcher:
            query_parser = QueryParser("content", schema=self.parent.ix.schema)
            query_parser.add_plugin(RegexPlugin())
            search_terms = self.search_term
            if "ignore_whitespace" in self.word_search_option:
                search_terms = [r'\s*'.join(term) for term in self.search_term.split()]
            if "words_together" in self.word_search_option:
                search_terms = ' AND '.join(self.search_term.split())
            query = query_parser.parse(search_terms)
            
            results = searcher.search(query, limit=None)
            total_results = len(results)
            
            for i, result in enumerate(results):
                if self.cancelled:
                    break
                self.results = {}
                title = result['title']
                path = result['path']
                paragraph_index = result['paragraph_number']
                if (title, path) not in self.results:
                    self.results[(title, path)] = []
                self.results[(title, path)].append(paragraph_index)
                self.progress_update.emit(i + 1, total_results)

        self.finished.emit()
    
    def cancel(self):
        self.cancelled = True

    def search_words_together(self, word_doc, search_term):
        # Search for search term in Word documents
        try:
            doc = Document(word_doc)
            body = doc._body._body
            ps = body.xpath(".//w:p")
            for p in ps:
                paragraph = Paragraph(p, None)
                p_nospace = paragraph.text.replace(" ", "").lower()
                if search_term.replace(" ", "").lower() in p_nospace:
                    return word_doc
        except Exception:
            return None
    
        return None
    
    def search_words_separately(self, word_doc, search_term):
        try:
            doc = Document(word_doc)
            body = doc._body._body
            ps = body.xpath(".//w:p")
            for p in ps:
                paragraph = Paragraph(p, None)
                words = search_term.split()
                p_nospace = paragraph.text.replace(" ", "").lower()
                if all(word.replace(" ", "").lower() in p_nospace for word in words):
                    return word_doc
        except Exception:
            return None

        return None

class ProgressDialog(QDialog):
    cancelled = pyqtSignal()

    def __init__(self, parent):
        super().__init__(parent)
        self.setWindowTitle('Searching...')
        self.resize(300, 100)

        self.progress_bar = QProgressBar()
        self.progress_label = QLabel()
        self.cancel_button = QPushButton('Cancel')
        self.cancel_button.clicked.connect(self.cancel)
        # self.progress_bar.setMaximum(num_files)
        self.progress_bar.setValue(0)
        # self.num_files = num_files

        vbox = QVBoxLayout()
        vbox.addWidget(self.progress_bar)
        vbox.addWidget(self.progress_label)
        vbox.addWidget(self.cancel_button)
        self.setLayout(vbox)
        # self.layout = vbox

    @pyqtSlot(int, int)
    def update_progress(self, processed_files, num_files):
        self.progress_bar.setMaximum(num_files)
        self.progress_bar.setValue(processed_files)
        self.progress_label.setText(f'{processed_files} out of {num_files} files processed')

    def cancel(self):
        self.cancelled.emit()

class DocumentItartor:
    def __init__(self, document, paragraph_indices, range=1) -> None:
        with open(document, 'rb') as f:
            document = Document(f)
            body = document._body._body
            self.paragraphs = body.xpath(".//w:p")

        self.found_indices = paragraph_indices
        self.current_index = 0
        self.range = range
        self.previous_paragraph = None
        self.current_paragraph = self.get_paragraph_at(0)
        self.next_paragraph = self.get_paragraph_at(1)
    
    def move_next(self):
        self.current_index += 1
        self.previous_paragraph = self.current_paragraph
        self.current_paragraph = self.next_paragraph
        self.next_paragraph = self.get_paragraph_at(self.current_index + 1)

    def move_previous(self):
        self.current_index -= 1
        self.next_paragraph = self.current_paragraph
        self.current_paragraph = self.previous_paragraph
        self.previous_paragraph = self.get_paragraph_at(self.current_index - 1)

    def get_paragraph_at(self, index):
        if 0 <= index < len(self.found_indices):
            return self.get_inrange_paragraphs(self.found_indices[index])
        return None
    
    def get_inrange_paragraphs(self, index):
        if index < 0 or index >= len(self.paragraphs):
            return None
        start = max(0, index - self.range)
        end = min(len(self.paragraphs), index + self.range + 1)
        sublist = self.paragraphs[start:end]
        sublist = map(lambda p: "        " + Paragraph(p, None).text, sublist)

        return "\n".join(sublist)

    def get_index(self):
        return self.found_indices[self.current_index] if 0 <= self.current_index < len(self.found_indices) else -1

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Set window title and size
        self.setWindowTitle('docxearch - Word Document Search')
        self.setGeometry(100, 100, 800, 600)
        font = QFont("Arial", 13)
        font.setWeight(QFont.Light)
        self.setFont(font)

        # Create central widget and layout
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Save last directory searched, on start is root
        self.last_directory = '/'

        # Create search term input
        search_layout = QHBoxLayout()
        left_layout = QVBoxLayout()
        right_layout = QVBoxLayout()
        self.label = QLabel('Search for:')
        # self.label.setFixedHeight(50)
        self.textbox = QLineEdit()
        left_h_layout_top = QHBoxLayout()
        left_h_layout_top.addWidget(self.label)
        left_layout.addLayout(left_h_layout_top)
        left_h_layout = QHBoxLayout()
        dummy_label = QLabel('         ')
        left_h_layout.addWidget(dummy_label)
        left_layout.addLayout(left_h_layout)
        search_layout.addLayout(left_layout)
        right_h_layout = QHBoxLayout()
        right_h_layout.addWidget(self.textbox)

        # Create search button
        self.search_button = QPushButton('Search')
        self.search_button.clicked.connect(self.search_word_docs)
        right_h_layout.addWidget(self.search_button)
        right_layout.addLayout(right_h_layout)

        # Add word checkbox
        options_layout = QHBoxLayout()
        self.word_checkbox = QCheckBox(self.central_widget)
        self.word_checkbox.setGeometry(QRect(610, 40, 151, 31))
        self.word_checkbox.setObjectName("word_checkbox")
        self.word_checkbox.setText("search by word")
        self.word_checkbox.stateChanged.connect(self.update_word_search_option)
        options_layout.addWidget(self.word_checkbox)
        self.search_again_checkbox = QCheckBox()
        self.search_again_checkbox.setGeometry(QRect(610, 40, 151, 31))
        self.search_again_checkbox.setObjectName("search_again_checkbox")
        self.search_again_checkbox.setText("search again from list")
        self.search_again_checkbox.stateChanged.connect(self.update_search_again_option)
        options_layout.addWidget(self.search_again_checkbox)

        right_layout.addLayout(options_layout)
        search_layout.addLayout(right_layout)
        self.layout.addLayout(search_layout)

        # Create results table
        self.table = QTableWidget()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(['File name', 'Path'])
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.table.cellPressed.connect(self.show_single_doc_search)
        self.table.cellDoubleClicked.connect(self.open_doc)
        self.layout.addWidget(self.table)

        self.current_row = None

         # Create results label
        self.results_label = QLabel('')
        self.layout.addWidget(self.results_label)

        # Default options to False
        self.word_search_option = {"same_paragraph", "ignore_whitespace"}
        self.search_again_option = False

        # Default variables
        self.file_list = []

        # Create text box
        self.text_browser = QTextBrowser()
        font = self.text_browser.document().defaultFont()
        font.setFamily("Gothic")
        font.setPointSize(13)
        self.text_browser.document().setDefaultFont(font)

        # Create previous and next buttons below text box
        self.prev_button = QPushButton('Previous')
        self.prev_button.setEnabled(False)
        self.next_button = QPushButton('Next')
        
        self.doc_text_nav = QWidget()
        vbox = QVBoxLayout()
        vbox.addWidget(self.text_browser)
        hbox = QHBoxLayout()
        hbox.addWidget(self.prev_button)
        hbox.addWidget(self.next_button)
        hbox.addStretch(1)
        vbox.addLayout(hbox)
        self.doc_text_nav.setLayout(vbox)
        self.doc_text_nav.setVisible(False)
        self.layout.addWidget(self.doc_text_nav)
        
        self.prev_button.clicked.connect(self.prev_paragraph)
        self.next_button.clicked.connect(self.next_paragraph)
    
    def initialize_whoosh_index(self):
        schema = Schema(title=TEXT(stored=True), path=ID(stored=True), content=TEXT, paragraph_number=NUMERIC(stored=True))
        if not os.path.exists("indexdir"):
            os.mkdir("indexdir")
            self.ix = create_in("indexdir", schema)
        else:
            self.ix = open_dir("indexdir")

    def update_whoosh_index(self, file_list):
        writer = self.ix.writer()
        for file_path in file_list:
            try:
                doc = Document(file_path)
                for i, paragraph in enumerate(doc.paragraphs):
                    content = paragraph.text
                    writer.add_document(
                        title=os.path.basename(file_path),
                        path=file_path,
                        content=content,
                        paragraph_number=i
                    )
            except Exception as e:
                print(f"Error indexing {file_path}: {e}")
        writer.commit()

    def search_word_docs(self):
        # Get search term from textbox
        self.search_term = self.textbox.text().strip()
        if not self.search_term:
            message_box = QMessageBox()
            message_box.setWindowTitle("Error")
            message_box.setText("Search term can't be empty!")
            message_box.setStandardButtons(QMessageBox.Ok)
            message_box.setDefaultButton(QMessageBox.Ok)
            message_box.exec_()
            return

        # Set loading indicator
        self.search_button.setDisabled(True)
        self.doc_text_nav.setVisible(False)

        thread = QThread(self)

        worker = SearchWordsThread(self, self.search_term, self.word_search_option)
        progress_dialog = ProgressDialog(self)
        worker.moveToThread(thread)
        worker.progress_update.connect(progress_dialog.update_progress)
        # worker.progress_update.connect(lambda processed, total: QMetaObject.invokeMethod(progress_dialog, 'update_progress', Qt.QueuedConnection, Q_ARG(int, processed), Q_ARG(int, total)))
        
        # Wait for worker to finish
        worker.finished.connect(lambda: self.show_results(worker.results))
        worker.finished.connect(progress_dialog.close)
        worker.finished.connect(thread.quit)
        progress_dialog.cancelled.connect(lambda: self.cancel_search(worker))
        progress_dialog.cancelled.connect(thread.quit)
        progress_dialog.cancelled.connect(progress_dialog.close)

        if not self.search_again_option:
            # Open directory dialog to select directory to search
            dir_path = QFileDialog.getExistingDirectory(self, 'Open Directory', self.last_directory)
            print("dir path: ", dir_path)
            locator = FindAllDocXFiles(self, dir_path)
            locator.moveToThread(thread)
            locator.progress_update.connect(progress_dialog.update_progress)
            locator.finished.connect(self.update_whoosh_index)
            progress_dialog.cancelled.connect(lambda: self.cancel_search(locator))
            thread.started.connect(locator.run)
            # Set last directory to this search
            self.last_directory = dir_path
        else:
            thread.started.connect(worker.run)

        thread.start()
        progress_dialog.exec_()

    def cancel_search(self, worker: FindAllDocXFiles | SearchWordsThread):
        worker.cancel()
        worker.quit()
        self.search_button.setEnabled(True)
    
    def show_results(self, results):
        # Display results in table
        self.table.setRowCount(len(results))
        for i, ((file_name, path), paragraph_index) in enumerate(results):
            self.table.setItem(i, 0, QTableWidgetItem(file_name))
            self.table.setItem(i, 1, QTableWidgetItem(path))
            self.table.item(i, 0).setData(Qt.ItemDataRole.UserRole, paragraph_index)
        
        for row in range(self.table.rowCount()):
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
        
        self.results_label.setText(f'Returned {len(results)} results.')

        self.search_button.setEnabled(True)

        # Save results
        self.file_list = [result[0][1] for result in results]
    
    def show_single_doc_search(self, row, col):
        # don't start from beginning if already in the row
        if self.current_row and self.current_row == row:
            return
        self.current_row = row

        file_path = self.table.item(row, 1).text()
        paragraph_indices = self.table.item(row,0).data(Qt.UserRole)
        self.doc_paragraphs = DocumentItartor(file_path, paragraph_indices)
        self.display_paragraph()
        self.doc_text_nav.setVisible(True)

    def display_paragraph(self):
        # Clear the text browser
        self.text_browser.clear()
        # Set the font family and size
        font = QFont("Nanum Gothic", 11)
        # self.text_browser.setFont(font)
        if not self.doc_paragraphs.previous_paragraph:
            self.prev_button.setEnabled(False)
        else:
            self.prev_button.setEnabled(True)
        if not self.doc_paragraphs.next_paragraph:
            self.next_button.setEnabled(False)
        else:
            self.next_button.setEnabled(True)

        paragraph = self.doc_paragraphs.current_paragraph
        paragraph_without_spaces = paragraph.replace(" ", "")
        whitespace_counts = []
        count = 0
        for i in range(len(paragraph)):
            if paragraph[i] == " ":
                count += 1
            else:
                whitespace_counts.append(count)

        doc = QTextDocument(paragraph)
        doc.setDefaultFont(font)
        
        # Highlight the search terms in the paragraph
        format = QTextCharFormat()
        format.setBackground(QBrush(QColor("yellow")))
        text_cursor = QTextCursor(doc)
        if self.word_search_option:
            for term in self.search_term.split():
                regex = QRegularExpression(term, QRegularExpression.CaseInsensitiveOption)
                matches = regex.globalMatch(paragraph_without_spaces)
                while matches.hasNext():
                    match = matches.next()
                    if match.isValid:
                        start_index = match.capturedStart() + whitespace_counts[match.capturedStart()]
                        end_index = match.capturedEnd() + whitespace_counts[match.capturedEnd()]
                        text_cursor.setPosition(start_index)
                        text_cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, end_index - start_index)
                        text_cursor.mergeCharFormat(format)
        else:
            pattern = self.search_term.replace(" ", "")
            regex = QRegularExpression(pattern, QRegularExpression.CaseInsensitiveOption)
            matches = regex.globalMatch(paragraph_without_spaces)
            while matches.hasNext():
                match = matches.next()
                if match.isValid:
                    start_index = match.capturedStart() + whitespace_counts[match.capturedStart()]
                    end_index = match.capturedEnd() + whitespace_counts[match.capturedEnd()]
                    text_cursor.setPosition(start_index)
                    text_cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, end_index - start_index)
                    text_cursor.mergeCharFormat(format)
        
        self.text_browser.setDocument(doc)
    
    def prev_paragraph(self):
        self.doc_paragraphs.move_previous()
        self.display_paragraph()
    
    def next_paragraph(self):
        self.doc_paragraphs.move_next()
        self.display_paragraph()

    def open_doc(self, row, col):
        file_path = self.table.item(row, 1).text()

        try:
            # Open Word document and highlight search term
            word_doc = win32.gencache.EnsureDispatch('Word.Application')
            doc = word_doc.Documents.Open(r'"{}"'.format(file_path))
            index = self.doc_paragraphs.get_index() + 1
            if index > 0:
                paragraph = doc.Paragraphs(index)
            else:
                raise Exception(f"Index out of range: {index}")

            # Select the paragraph
            paragraph.Range.Select()
            word_doc.Visible = True
            word_doc.Activate()

        except Exception as ex:
            print('Error while opening doc.')
            message_box = QMessageBox()
            message_box.setWindowTitle("Error")
            message_box.setText("Error while trying to open Word doc.: {}".format(ex))
            message_box.setStandardButtons(QMessageBox.Ok)
            message_box.setDefaultButton(QMessageBox.Ok)
            message_box.exec_()

    def update_word_search_option(self):
        if self.word_checkbox.isChecked():
            self.word_search_option.add("words_together") 
        else:
            self.word_search_option.remove("words_together")

    def update_search_again_option(self):
        self.search_again_option = self.search_again_checkbox.isChecked()
    
    def setup_ui(self):
        pass

if __name__ == '__main__':
    # Create the application
    app = QApplication([])

    # Create the main window
    window = MainWindow()
    window.initialize_whoosh_index()

    # Set up the user interface
    window.setup_ui()

    # Show the window
    window.show()

    # Run the event loop
    app.exec_()
