from PyQt5.QtWidgets import QApplication, QMainWindow, QWidget, QTableWidget, QTableWidgetItem, QVBoxLayout, QHBoxLayout, QLabel, QLineEdit, QPushButton, QFileDialog, QCheckBox, QMessageBox, QHeaderView, QProgressBar, QDialog, QTextBrowser
from PyQt5.QtGui import QTextCursor, QFont, QTextCharFormat, QColor, QBrush, QTextDocument
from PyQt5.QtCore import Qt, QRect, QThread, pyqtSignal, QRegularExpression
from docx import Document
from docx.text.paragraph import Paragraph
import os
import win32com.client as win32
import re
import concurrent.futures

class SearchWordsThread(QThread):
    progress_update = pyqtSignal(int, int)
    finished = pyqtSignal()

    def __init__(self, parent, dir_path, search_term, word_search_option):
        super().__init__()
        # self.queue = queue
        self.parent = parent
        self.dir_path = dir_path
        self.search_term = search_term
        self.word_search_option = word_search_option
        self.cancelled = False
        self.executor = concurrent.futures.ThreadPoolExecutor()
        self.results = []
    
    def run(self):
        # Search for Word documents in directory
        word_docs = []
        num_files = 0
        try:
            for root, dirs, files in os.walk(self.dir_path):
                for file in files:
                    if file.endswith('.docx'):
                        word_docs.append(os.path.join(root, file))
                        num_files += 1
                        self.progress_update.emit(0, num_files)
        except Exception as e:
            message_box = QMessageBox(self.parent)
            message_box.setWindowTitle("Error while finding docx files.")
            message_box.setText(e)
            message_box.setStandardButtons(QMessageBox.Ok)
            message_box.setDefaultButton(QMessageBox.Ok)
            message_box.exec_()
            self.finished.emit()
            return

        num_processed = 0

        futures = []
        for word_doc in word_docs:
            if self.cancelled:
                break
            if self.word_search_option:
                future = self.executor.submit(self.search_words_separately, word_doc, self.search_term)
            else:
                future = self.executor.submit(self.search_words_together, word_doc, self.search_term)
            futures.append(future)
        
        for future in concurrent.futures.as_completed(futures):
            if self.cancelled:
                break
            result = future.result()
            if result is not None:
                self.results.append((os.path.basename(result), result.replace('/', '\\')))
            num_processed += 1
            self.progress_update.emit(num_processed, num_files)
    
        self.executor.shutdown(wait=True, cancel_futures=True)
        self.finished.emit()
    
    def cancel(self):
        self.cancelled = True

    def search_words_together(self, word_doc, search_term):
        # Search for search term in Word documents
        try:
            doc = Document(word_doc)
            body = doc._body._body
            ps = body.xpath(".//w:p")
            for p in ps:
                paragraph = Paragraph(p, None)
                p_nospace = paragraph.text.replace(" ", "").lower()
                if search_term.replace(" ", "").lower() in p_nospace:
                    return word_doc
        except Exception:
            return None
    
        return None
    
    def search_words_separately(self, word_doc, search_term):
        try:
            doc = Document(word_doc)
            body = doc._body._body
            ps = body.xpath(".//w:p")
            for p in ps:
                paragraph = Paragraph(p, None)
                words = search_term.split()
                p_nospace = paragraph.text.replace(" ", "").lower()
                if all(word.replace(" ", "").lower() in p_nospace for word in words):
                    return word_doc
        except Exception:
            return None

        return None

class ProgressDialog(QDialog):
    cancelled = pyqtSignal()

    def __init__(self, parent):
        super().__init__(parent)
        self.setWindowTitle('Searching...')
        self.resize(300, 100)

        self.progress_bar = QProgressBar()
        self.progress_label = QLabel()
        self.cancel_button = QPushButton('Cancel')
        self.cancel_button.clicked.connect(self.cancel)
        # self.progress_bar.setMaximum(num_files)
        self.progress_bar.setValue(0)
        # self.num_files = num_files

        vbox = QVBoxLayout()
        vbox.addWidget(self.progress_bar)
        vbox.addWidget(self.progress_label)
        vbox.addWidget(self.cancel_button)
        self.setLayout(vbox)
        # self.layout = vbox

    def update_progress(self, processed_files, num_files):
        self.progress_bar.setMaximum(num_files)
        self.progress_bar.setValue(processed_files)
        self.progress_label.setText(f'{processed_files} out of {num_files} files processed')

    def cancel(self):
        self.cancelled.emit()

class DocumentItartor:
    def __init__(self, document, search_terms, search_by_word, range=1) -> None:
        with open(document, 'rb') as f:
            document = Document(f)
            body = document._body._body
            self.paragraphs = body.xpath(".//w:p")

        self.found_indices = []
        self.search_terms = search_terms
        self.search_by_word = search_by_word
        self.current_index = 0
        self.iterator = self.find_next()
        self.range = range
        self.previous_paragraph = None
        self.current_paragraph = self.get_paragraph_at(0)
        self.next_paragraph = self.get_paragraph_at(1)
    
    def move_next(self):
        self.current_index += 1
        self.previous_paragraph = self.current_paragraph
        self.current_paragraph = self.next_paragraph
        self.next_paragraph = self.get_paragraph_at(self.current_index + 1)

    def move_previous(self):
        self.current_index -= 1
        self.next_paragraph = self.current_paragraph
        self.current_paragraph = self.previous_paragraph
        self.previous_paragraph = self.get_paragraph_at(self.current_index - 1)

    def get_paragraph_at(self, index):
        if 0 <= index < len(self.found_indices):
            return self.get_inrange_paragraphs(self.found_indices[index])
        if index >= len(self.found_indices):
            try:
                i = next(self.iterator)
            except StopIteration:
                return None
            self.found_indices.append(i)
            return self.get_inrange_paragraphs(i)
        return None
    
    def find_next(self):
        for i in range(len(self.paragraphs)):
            p = self.paragraphs[i]
            paragraph = Paragraph(p, None)
            paragraph_without_spaces = paragraph.text.replace(" ", "").lower()
            if self.search_by_word:
                words = self.search_terms.split()
                if all(word.lower() in paragraph_without_spaces for word in words):
                    yield i
            else:
                pattern = self.search_terms.replace(" ", "").lower()
                match = re.search(pattern, paragraph_without_spaces)
                if match:
                    yield i
    
    def get_inrange_paragraphs(self, index):
        if index < 0 or index >= len(self.paragraphs):
            return None
        start = max(0, index - self.range)
        end = min(len(self.paragraphs), index + self.range + 1)
        sublist = self.paragraphs[start:end]
        sublist = map(lambda p: "        " + Paragraph(p, None).text, sublist)

        return "\n".join(sublist)

    def get_index(self):
        return self.found_indices[self.current_index] if 0 <= self.current_index < len(self.found_indices) else -1

class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()

        # Set window title and size
        self.setWindowTitle('docxearch - Word Document Search')
        self.setGeometry(100, 100, 800, 600)
        font = QFont("Arial", 13)
        font.setWeight(QFont.Light)
        self.setFont(font)

        # Create central widget and layout
        self.central_widget = QWidget()
        self.setCentralWidget(self.central_widget)
        self.layout = QVBoxLayout()
        self.central_widget.setLayout(self.layout)

        # Save last directory searched, on start is root
        self.last_directory = '/'

        # Create search term input
        search_layout = QHBoxLayout()
        left_layout = QVBoxLayout()
        right_layout = QVBoxLayout()
        self.label = QLabel('Search for:')
        # self.label.setFixedHeight(50)
        self.textbox = QLineEdit()
        left_h_layout_top = QHBoxLayout()
        left_h_layout_top.addWidget(self.label)
        left_layout.addLayout(left_h_layout_top)
        left_h_layout = QHBoxLayout()
        dummy_label = QLabel('         ')
        left_h_layout.addWidget(dummy_label)
        left_layout.addLayout(left_h_layout)
        search_layout.addLayout(left_layout)
        right_h_layout = QHBoxLayout()
        right_h_layout.addWidget(self.textbox)

        # Create search button
        self.search_button = QPushButton('Search')
        self.search_button.clicked.connect(self.search_word_docs)
        right_h_layout.addWidget(self.search_button)
        right_layout.addLayout(right_h_layout)

        # Add word checkbox
        options_layout = QHBoxLayout()
        self.word_checkbox = QCheckBox(self.central_widget)
        self.word_checkbox.setGeometry(QRect(610, 40, 151, 31))
        self.word_checkbox.setObjectName("word_checkbox")
        self.word_checkbox.setText("search by word")
        self.word_checkbox.stateChanged.connect(self.update_word_search_option)
        options_layout.addWidget(self.word_checkbox)

        right_layout.addLayout(options_layout)
        search_layout.addLayout(right_layout)
        self.layout.addLayout(search_layout)

        # Create results table
        self.table = QTableWidget()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(['File name', 'Path'])
        header = self.table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.ResizeMode.ResizeToContents)
        header.setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.table.cellPressed.connect(self.show_single_doc_search)
        self.table.cellDoubleClicked.connect(self.open_doc)
        self.layout.addWidget(self.table)

        self.current_row = None

         # Create results label
        self.results_label = QLabel('')
        self.layout.addWidget(self.results_label)

        # Default options to False
        self.word_search_option = False
        self.highlight_option = False

        # Create text box
        self.text_browser = QTextBrowser()
        font = self.text_browser.document().defaultFont()
        font.setFamily("Gothic")
        font.setPointSize(13)
        self.text_browser.document().setDefaultFont(font)

        # Create previous and next buttons below text box
        self.prev_button = QPushButton('Previous')
        self.prev_button.setEnabled(False)
        self.next_button = QPushButton('Next')
        
        self.doc_text_nav = QWidget()
        vbox = QVBoxLayout()
        vbox.addWidget(self.text_browser)
        hbox = QHBoxLayout()
        hbox.addWidget(self.prev_button)
        hbox.addWidget(self.next_button)
        hbox.addStretch(1)
        vbox.addLayout(hbox)
        self.doc_text_nav.setLayout(vbox)
        self.doc_text_nav.setVisible(False)
        self.layout.addWidget(self.doc_text_nav)
        
        self.prev_button.clicked.connect(self.prev_paragraph)
        self.next_button.clicked.connect(self.next_paragraph)

    def search_word_docs(self):
        # Get search term from textbox
        self.search_term = self.textbox.text().strip()
        if not self.search_term:
            message_box = QMessageBox()
            message_box.setWindowTitle("Error")
            message_box.setText("Search term can't be empty!")
            message_box.setStandardButtons(QMessageBox.Ok)
            message_box.setDefaultButton(QMessageBox.Ok)
            message_box.exec_()
            return

        # Open directory dialog to select directory to search
        dir_path = QFileDialog.getExistingDirectory(self, 'Open Directory', self.last_directory)

        # Set loading indicator
        self.search_button.setDisabled(True)
        self.doc_text_nav.setVisible(False)

        thread = QThread(self)
        worker = SearchWordsThread(self, dir_path, self.search_term, self.word_search_option)
        progress_dialog = ProgressDialog(self)
        worker.moveToThread(thread)
        worker.progress_update.connect(progress_dialog.update_progress)
        
        # Wait for worker to finish
        worker.finished.connect(lambda: self.show_results(worker.results))
        worker.finished.connect(progress_dialog.close)
        progress_dialog.cancelled.connect(thread.quit)
        progress_dialog.cancelled.connect(lambda: self.cancel_search(worker))

        thread.started.connect(worker.run)
        thread.start()
        progress_dialog.exec_()

        # Set last directory to this search
        self.last_directory = dir_path
    
    def cancel_search(self, worker: SearchWordsThread):
        worker.cancel()
        worker.quit()
    
    def show_results(self, results):
        # Display results in table
        self.table.setRowCount(len(results))
        for i, (file_name, path) in enumerate(results):
            self.table.setItem(i, 0, QTableWidgetItem(file_name))
            self.table.setItem(i, 1, QTableWidgetItem(path))
        
        for row in range(self.table.rowCount()):
            for col in range(self.table.columnCount()):
                item = self.table.item(row, col)
                item.setFlags(item.flags() & ~Qt.ItemIsEditable)
        
        self.results_label.setText(f'Returned {len(results)} results.')

        self.search_button.setEnabled(True)
    
    def show_single_doc_search(self, row, col):
        # don't start from beginning if already in the row
        if self.current_row and self.current_row == row:
            return
        self.current_row = row

        file_path = self.table.item(row, 1).text()
        self.doc_paragraphs = DocumentItartor(file_path, self.search_term, self.word_search_option)
        self.display_paragraph()
        self.doc_text_nav.setVisible(True)

    def display_paragraph(self):
        # Clear the text browser
        self.text_browser.clear()
        # Set the font family and size
        font = QFont("Nanum Gothic", 11)
        # self.text_browser.setFont(font)
        if not self.doc_paragraphs.previous_paragraph:
            self.prev_button.setEnabled(False)
        else:
            self.prev_button.setEnabled(True)
        if not self.doc_paragraphs.next_paragraph:
            self.next_button.setEnabled(False)
        else:
            self.next_button.setEnabled(True)

        paragraph = self.doc_paragraphs.current_paragraph
        paragraph_without_spaces = paragraph.replace(" ", "")
        whitespace_counts = []
        count = 0
        for i in range(len(paragraph)):
            if paragraph[i] == " ":
                count += 1
            else:
                whitespace_counts.append(count)

        doc = QTextDocument(paragraph)
        doc.setDefaultFont(font)
        
        # Highlight the search terms in the paragraph
        format = QTextCharFormat()
        format.setBackground(QBrush(QColor("yellow")))
        text_cursor = QTextCursor(doc)
        if self.word_search_option:
            for term in self.search_term.split():
                regex = QRegularExpression(term, QRegularExpression.CaseInsensitiveOption)
                matches = regex.globalMatch(paragraph_without_spaces)
                while matches.hasNext():
                    match = matches.next()
                    if match.isValid:
                        start_index = match.capturedStart() + whitespace_counts[match.capturedStart()]
                        end_index = match.capturedEnd() + whitespace_counts[match.capturedEnd()]
                        text_cursor.setPosition(start_index)
                        text_cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, end_index - start_index)
                        text_cursor.mergeCharFormat(format)
        else:
            pattern = self.search_term.replace(" ", "")
            regex = QRegularExpression(pattern, QRegularExpression.CaseInsensitiveOption)
            matches = regex.globalMatch(paragraph_without_spaces)
            while matches.hasNext():
                match = matches.next()
                if match.isValid:
                    start_index = match.capturedStart() + whitespace_counts[match.capturedStart()]
                    end_index = match.capturedEnd() + whitespace_counts[match.capturedEnd()]
                    text_cursor.setPosition(start_index)
                    text_cursor.movePosition(QTextCursor.Right, QTextCursor.KeepAnchor, end_index - start_index)
                    text_cursor.mergeCharFormat(format)
        
        self.text_browser.setDocument(doc)
    
    def prev_paragraph(self):
        self.doc_paragraphs.move_previous()
        self.display_paragraph()
    
    def next_paragraph(self):
        self.doc_paragraphs.move_next()
        self.display_paragraph()

    def open_doc(self, row, col):
        # Get search term from textbox
        # search_term = self.textbox.text()

        file_path = self.table.item(row, 1).text()
        # print(file_path)

        # word.Visible = True
        try:
            # Open Word document and highlight search term
            word_doc = win32.gencache.EnsureDispatch('Word.Application')
            doc = word_doc.Documents.Open(r'"{}"'.format(file_path))
            index = self.doc_paragraphs.get_index() + 1
            if index > 0:
                paragraph = doc.Paragraphs(index)
            else:
                raise Exception(f"Index out of range: {index}")

            # Select the paragraph
            paragraph.Range.Select()
            word_doc.Visible = True
            word_doc.Activate()

            # doc_content = doc.Content
            # first_occurence = True

            # for p in doc_content.Paragraphs:
            #     words = search_term.split()
            #     paragraph_text = p.Range.Text.strip()

            #     if self.word_search_option:
            #         if all(word in paragraph_text for word in words):
            #             # Highlight all occurrences of the words within the paragraph
            #             search_range = p.Range
            #             search_range.Find.ClearFormatting()
            #             search_range.Find.Forward = True
            #             if first_occurence:
            #                 search_range.Find.Execute(words[0], MatchWholeWord=False, MatchCase=False, MatchWildcards=False)
            #                 search_range.Select()
            #                 first_occurence = False

            #             if self.highlight_option:
            #                 search_range = p.Range
            #                 search_range.Find.ClearFormatting()
            #                 for word in words:
            #                     search_range.Find.Replacement.Highlight = True
            #                     search_range.Find.Execute(word, Replace=2, MatchWholeWord=False, MatchCase=False, MatchWildcards=False)
                
            #     else:
            #         pattern = re.compile("\s*".join(words))
            #         match = re.search(pattern, paragraph_text)
            #         if match:
            #             search_range = p.Range
            #             search_range.Find.ClearFormatting()
            #             search_range.Find.Execute(match.group(), MatchWholeWord=False, MatchCase=False, MatchWildcards=False)
            #             if first_occurence:
            #                 search_range.Select()
            #                 first_occurence = False
            #             if self.highlight_option:
            #                 search_range.HighlightColorIndex = win32.constants.wdYellow
            #             else:
            #                 break

        except Exception as ex:
            print('Error while opening doc.')
            message_box = QMessageBox()
            message_box.setWindowTitle("Error")
            message_box.setText("Error while trying to open Word doc.: {}".format(ex))
            message_box.setStandardButtons(QMessageBox.Ok)
            message_box.setDefaultButton(QMessageBox.Ok)
            message_box.exec_()

    def update_word_search_option(self):
        self.word_search_option = self.word_checkbox.isChecked()

    def update_highlight_option(self):
        self.highlight_option = self.highlight_checkbox.isChecked()
    
    def setup_ui(self):
        pass

if __name__ == '__main__':
    # Create the application
    app = QApplication([])

    # Create the main window
    window = MainWindow()

    # Set up the user interface
    window.setup_ui()

    # Show the window
    window.show()

    # Run the event loop
    app.exec_()
